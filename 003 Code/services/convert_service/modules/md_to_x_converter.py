from typing import Callable, Dict, cast
from markdown2 import markdown
import pdfkit  
import pandas as pd
from io import BytesIO, StringIO
from docx import Document
import pandas as pd 

def convert_markdown_to_format(md_content: str, fmt: str) -> bytes | None:
    handlers: Dict[str, Callable[[str], bytes]] = {
        "html": to_html,
        "pdf": to_pdf,
        "csv": to_csv,
        "xlsx": to_xlsx,
        "json": to_json,
        "docx": to_docx
    }

    return handlers.get(fmt, lambda _: None)(md_content)

def to_html(content: str) -> bytes:
    return f"""
        <html>
            <head>
                <meta charset="utf-8">
                <style>
                    body {{ font-family: "Apple SD Gothic Neo", "Nanum Gothic", sans-serif; line-height: 1.6; padding: 20px; font-size: 12pt; }}
                    table {{ border-collapse: collapse; width: 100%; margin: 20px 0; font-size: 10pt; }}
                    th, td {{ border: 1px solid #444; padding: 8px 12px; text-align: center; }}
                    th {{ background-color: #f2f2f2; font-weight: bold; }}
                    tr:nth-child(even) {{ background-color: #fafafa; }}
                </style>
            </head>
            <body>
                {markdown(content, extras=["tables"])}
            </body>
        </html>
    """.encode("utf-8")

def to_pdf(content: str) -> bytes:
    html = to_html(content).decode("utf-8")
    pdf = pdfkit.from_string(html, False)
    return cast(bytes, pdf)

def to_csv(content: str) -> bytes:
    # df = to_dataframe(content)
    df = pd.read_csv(StringIO(content))

    buffer = BytesIO()
    df.to_csv(buffer, index=False, encoding="utf-8")
    return buffer.getvalue()

def to_xlsx(content: str) -> bytes:
    # df = to_dataframe(content)
    df = pd.read_csv(StringIO(content))

    buffer = BytesIO()
    df.to_excel(buffer, index=False)
    return buffer.getvalue()

def to_json(content: str) -> bytes:
    # df = to_dataframe(content)
    df = pd.read_csv(StringIO(content))
    return df.to_json(orient="records", force_ascii=False, indent=2).encode("utf-8")

def to_dataframe(content: str) -> pd.DataFrame:
    lines = [line for line in content.splitlines() if "|" in line and "---" not in line]
    if len(lines) < 2:
        raise ValueError("유효한 마크다운 테이블이 아닙니다.")
    
    table = "\n".join(lines)
    df = pd.read_csv(StringIO(table), sep="|", engine="python", skipinitialspace=True)
    df = df.dropna(axis=1, how="all")  # 빈 열 제거

    return df

def to_docx(content: str) -> bytes:
    document = Document()
    lines = [line.strip() for line in content.splitlines() if "|" in line and "---" not in line]
    if lines:
        # 첫 줄은 헤더
        header = [h.strip() for h in lines[0].split("|") if h.strip()]
        rows = [[c.strip() for c in row.split("|") if c.strip()] for row in lines[1:]]

        # 표 생성
        table = document.add_table(rows=1, cols=len(header))
        table.style = "Table Grid"

        # 헤더 채우기
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(header):
            hdr_cells[i].text = h

        # 데이터 채우기
        for row in rows:
            row_cells = table.add_row().cells
            for i, c in enumerate(row):
                row_cells[i].text = c

        document.add_paragraph("\n")  # 표 이후 여백

    # 표 외 텍스트도 출력
    for line in content.splitlines():
        if "|" not in line:
            document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()